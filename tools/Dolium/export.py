#╔══════════════════════════════════════════════════════════════════════════════
#║ ⛨⛨⛨⛨⛨⛨⛨⛨⛨⛨⛨⛨⛨⛨⛨⛨⛨⛨⛨⛨⛨
#║ ⛨⛨⛨⛨⛨⛨⛨⛨⛨⛨⛨⛨⛨
#║ ⛨⛨⛨⛨⛨⛨⛨⛨
#║ ⛨⛨⛨⛨⛨
#║ ⛨⛨⛨
#║ ⛨⛨
#║ ⛨
#║ ⛨    The Dolium / export.py
#║ ⛨
#╚══════════════════════════════════════════════════════════════════════════════

from __future__ import annotations

import json
import re
import subprocess
import sys
from pathlib import Path

from models import Idea, chamber_name, chamber_latin


# ── Helpers ───────────────────────────────────────────────────────────────────

def _slug(title: str) -> str:
    """Lowercase, hyphen-separated, filesystem-safe slug from a title."""
    s = title.lower().strip()
    s = re.sub(r"[^\w\s-]", "", s)
    s = re.sub(r"[\s_]+", "-", s)
    s = re.sub(r"-+", "-", s).strip("-")
    return s or "untitled"


SECTION_LABELS = [
    ("body",            "Body"),
    ("motivation",      "Motivation"),
    ("scope_in",        "Scope — Inside"),
    ("scope_out",       "Scope — Outside"),
    ("system_map",      "System Map"),
    ("dependencies",    "Dependencies"),
    ("build_sequence",  "Build Sequence"),
    ("open_questions",  "Open Questions"),
    ("aesthetic_notes", "Aesthetic Notes"),
    ("declaration",     "Declaration"),
]


def _populated_sections(idea: Idea) -> list[tuple[str, str]]:
    """Return (label, content) pairs for all non-empty fields."""
    result = []
    for attr, label in SECTION_LABELS:
        val = getattr(idea, attr, "").strip()
        if val:
            result.append((label, val))
    return result


# ── ExportEngine ──────────────────────────────────────────────────────────────

class ExportEngine:
    """
    Generates export files from a fully populated Idea.
    JSON is always the source of truth — all other formats derived from it.
    """

    def __init__(self, output_dir: Path) -> None:
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate_all(self, idea: Idea) -> dict[str, Path]:
        return self.generate(idea, ["wiz", "docx", "md", "txt", "json"])

    def generate(self, idea: Idea, formats: list[str]) -> dict[str, Path]:
        results = {}
        for fmt in formats:
            try:
                fn = getattr(self, f"_to_{fmt}", None)
                if fn:
                    results[fmt] = fn(idea)
            except Exception as e:
                # Log but don't crash — partial export is better than none
                print(f"[export] {fmt} failed: {e}", file=sys.stderr)
        return results

    # ── Format writers ────────────────────────────────────────────────────────

    def _to_json(self, idea: Idea) -> Path:
        path = self.output_dir / f"{idea.id}.json"
        path.write_text(
            json.dumps(idea.to_dict(), indent=2, ensure_ascii=False),
            encoding="utf-8",
        )
        return path

    def _to_txt(self, idea: Idea) -> Path:
        path   = self.output_dir / f"{_slug(idea.title)}.txt"
        lines  = []
        rule   = "─" * 60

        lines.append("THE DOLIUM — PARATUM PACKAGE")
        lines.append(rule)
        lines.append(f"Title:    {idea.title}")
        lines.append(f"Chamber:  {chamber_name(idea.chamber)}  ({chamber_latin(idea.chamber)})")
        lines.append(f"Created:  {idea.created_at[:10]}")
        if idea.declared_at:
            lines.append(f"Declared: {idea.declared_at[:10]}")
        if idea.tags:
            lines.append(f"Tags:     {', '.join(idea.tags)}")
        lines.append(rule)

        for label, content in _populated_sections(idea):
            lines.append(f"\n{label.upper()}")
            lines.append(rule)
            lines.append(content)

        path.write_text("\n".join(lines), encoding="utf-8")
        return path

    def _to_md(self, idea: Idea) -> Path:
        path  = self.output_dir / f"{_slug(idea.title)}.md"
        lines = []

        lines.append(f"# {idea.title}")
        lines.append("")
        lines.append(f"**Chamber:** {chamber_name(idea.chamber)} — *{chamber_latin(idea.chamber)}*  ")
        lines.append(f"**Created:** {idea.created_at[:10]}  ")
        if idea.declared_at:
            lines.append(f"**Declared:** {idea.declared_at[:10]}  ")
        if idea.tags:
            lines.append(f"**Tags:** {', '.join(idea.tags)}  ")
        lines.append("")
        lines.append("---")
        lines.append("")

        for label, content in _populated_sections(idea):
            lines.append(f"## {label}")
            lines.append("")
            lines.append(content)
            lines.append("")

        path.write_text("\n".join(lines), encoding="utf-8")
        return path

    def _to_docx(self, idea: Idea) -> Path:
        """Clean Word document. No Wizard styling. For AI upload."""
        from docx import Document as DocxDocument
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc  = DocxDocument()
        path = self.output_dir / f"{_slug(idea.title)}.docx"

        # Title
        title_para = doc.add_heading(idea.title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Metadata
        meta = doc.add_paragraph()
        meta.add_run(f"Chamber: {chamber_name(idea.chamber)} — {chamber_latin(idea.chamber)}\n").bold = False
        meta.add_run(f"Created: {idea.created_at[:10]}\n")
        if idea.declared_at:
            meta.add_run(f"Declared: {idea.declared_at[:10]}\n")
        if idea.tags:
            meta.add_run(f"Tags: {', '.join(idea.tags)}")

        doc.add_paragraph()  # spacer

        # Sections
        for label, content in _populated_sections(idea):
            doc.add_heading(label, level=1)
            doc.add_paragraph(content)

        doc.save(str(path))
        return path

    def _to_wiz(self, idea: Idea) -> Path:
        """
        Wizard-styled document via export_wiz.js Node subprocess.
        Falls back gracefully if Node is not available.
        """
        script = Path(__file__).parent / "export_wiz.js"
        if not script.exists():
            raise FileNotFoundError(
                "export_wiz.js not found. "
                "The .wiz export requires the Node.js script to be present."
            )

        slug     = _slug(idea.title)
        out_path = self.output_dir / f"{slug}.wiz"

        payload = json.dumps(idea.to_dict(), ensure_ascii=False)

        try:
            result = subprocess.run(
                ["node", str(script), str(out_path)],
                input=payload,
                capture_output=True,
                text=True,
                timeout=30,
            )
            if result.returncode != 0:
                raise RuntimeError(result.stderr.strip() or "Node script failed.")
        except FileNotFoundError:
            raise RuntimeError(
                "Node.js is not installed or not on PATH. "
                "The .wiz export requires Node.js."
            )

        return out_path
