"""
export.py — Dolium v2
ExportEngine: generates .md .txt .json .docx .wiz from a completed Idea.
.wiz via Node.js subprocess — degrades gracefully if Node unavailable.
.docx via python-docx — degrades gracefully if unavailable.
"""

from __future__ import annotations

import json
import subprocess
import logging
from datetime import datetime, timezone

def _now_str() -> str:
    return datetime.now(timezone.utc).strftime('%Y-%m-%d')
from pathlib import Path
from typing import Optional

from models import Idea

log = logging.getLogger(__name__)


class ExportEngine:

    def __init__(self, exports_dir: Path):
        self._dir = Path(exports_dir)
        self._dir.mkdir(parents=True, exist_ok=True)

    def export_all(self, idea: Idea) -> dict[str, Optional[Path]]:
        """
        Export idea to all available formats.
        Returns dict of format → Path (None if export failed).
        """
        slug = self._slug(idea)
        results: dict[str, Optional[Path]] = {}

        results["json"] = self._to_json(idea, slug)
        results["md"]   = self._to_md(idea, slug)
        results["txt"]  = self._to_txt(idea, slug)
        results["docx"] = self._to_docx(idea, slug)
        results["wiz"]  = self._to_wiz(idea, slug)

        return results

    def export_format(self, idea: Idea, fmt: str) -> Optional[Path]:
        slug = self._slug(idea)
        handlers = {
            "json": self._to_json,
            "md":   self._to_md,
            "txt":  self._to_txt,
            "docx": self._to_docx,
            "wiz":  self._to_wiz,
        }
        fn = handlers.get(fmt)
        if fn is None:
            return None
        return fn(idea, slug)

    # ── Formats ───────────────────────────────────────────────────────────────

    def _to_json(self, idea: Idea, slug: str) -> Optional[Path]:
        path = self._dir / f"{slug}.json"
        try:
            path.write_text(
                json.dumps(idea.to_dict(), indent=2, ensure_ascii=False),
                encoding="utf-8",
            )
            return path
        except OSError as e:
            log.warning("JSON export failed: %s", e)
            return None

    def _to_md(self, idea: Idea, slug: str) -> Optional[Path]:
        path = self._dir / f"{slug}.md"
        try:
            path.write_text(self._build_md(idea), encoding="utf-8")
            return path
        except OSError as e:
            log.warning("Markdown export failed: %s", e)
            return None

    def _to_txt(self, idea: Idea, slug: str) -> Optional[Path]:
        path = self._dir / f"{slug}.txt"
        try:
            path.write_text(self._build_txt(idea), encoding="utf-8")
            return path
        except OSError as e:
            log.warning("TXT export failed: %s", e)
            return None

    def _to_docx(self, idea: Idea, slug: str) -> Optional[Path]:
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            log.warning("python-docx not installed — skipping .docx export")
            return None

        path = self._dir / f"{slug}.docx"
        try:
            doc = Document()
            doc.add_heading(idea.title or "(untitled)", level=1)
            doc.add_paragraph(f"Declared: {_now_str()}")
            doc.add_paragraph(f"Chamber: {idea.chamber_name()}")
            doc.add_paragraph("")

            sections = self._ordered_sections(idea)
            for heading, text in sections:
                if text.strip():
                    doc.add_heading(heading, level=2)
                    doc.add_paragraph(text.strip())

            doc.save(str(path))
            return path
        except Exception as e:
            log.warning("DOCX export failed: %s", e)
            return None

    def _to_wiz(self, idea: Idea, slug: str) -> Optional[Path]:
        """
        .wiz via Node.js + docx npm package. Degrades gracefully if unavailable.
        Generates a JSON manifest for the Node script to consume.
        """
        node_script = Path(__file__).parent / "wiz_export.js"
        if not node_script.exists():
            log.info(".wiz export skipped — wiz_export.js not found")
            return None

        manifest_path = self._dir / f"{slug}_wiz_manifest.json"
        output_path   = self._dir / f"{slug}.wiz"

        try:
            manifest_path.write_text(
                json.dumps(idea.to_dict(), indent=2, ensure_ascii=False),
                encoding="utf-8",
            )
            result = subprocess.run(
                ["node", str(node_script), str(manifest_path), str(output_path)],
                capture_output=True,
                text=True,
                timeout=15,
            )
            manifest_path.unlink(missing_ok=True)

            if result.returncode == 0:
                return output_path
            else:
                log.warning(".wiz export failed: %s", result.stderr)
                return None
        except (FileNotFoundError, subprocess.TimeoutExpired, OSError) as e:
            log.warning(".wiz export failed: %s", e)
            return None

    # ── Text builders ─────────────────────────────────────────────────────────

    def _build_md(self, idea: Idea) -> str:
        lines = [
            f"# {idea.title or '(untitled)'}",
            f"*Declared {_now_str()} · {idea.chamber_name()}*",
            "",
        ]
        for heading, text in self._ordered_sections(idea):
            if text.strip():
                lines += [f"## {heading}", "", text.strip(), ""]

        if idea.tags:
            lines += ["## Tags", "", " ".join(f"#{t}" for t in idea.tags), ""]

        return "\n".join(lines)

    def _build_txt(self, idea: Idea) -> str:
        lines = [
            idea.title.upper() if idea.title else "(UNTITLED)",
            f"Declared {_now_str()}",
            f"{idea.chamber_name()}",
            "",
            "─" * 60,
            "",
        ]
        for heading, text in self._ordered_sections(idea):
            if text.strip():
                lines += [heading.upper(), "", text.strip(), "", "─" * 40, ""]

        return "\n".join(lines)

    def _ordered_sections(self, idea: Idea) -> list[tuple[str, str]]:
        return [
            ("Body",          idea.body),
            ("Motivation",    idea.motivation),
            ("Elaboration",   idea.elaboration),
            ("Obstacles",     idea.obstacles),
            ("First Step",    idea.first_step),
            ("Refined Form",  idea.refined_form),
            ("Open Problems", idea.open_problems),
            ("Next Actions",  idea.next_actions),
            ("Declaration",   idea.declaration),
            ("Summary",       idea.summary),
        ]

    # ── Utilities ─────────────────────────────────────────────────────────────

    @staticmethod
    def _slug(idea: Idea) -> str:
        import re
        title = idea.title.lower().strip() if idea.title else "idea"
        slug  = re.sub(r"[^a-z0-9]+", "_", title)[:40].strip("_")
        ts    = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        return f"{slug}_{ts}" if slug else f"idea_{ts}"
